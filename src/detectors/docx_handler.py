import re
from docx import Document
from .hybrid_detector import hybrid_detector
from redaction.replacer import replacer
from data_structure import pii_type


class docx_handler:
    def __init__(self, input_path: str, output_path: str):
        self.input_path = input_path
        self.output_path = output_path
        self.detector = hybrid_detector()
        self.replacer = replacer()

        self.company_pattern = re.compile(
            r'\b(?:Limited|Ltd|Pvt|LLP|Inc|Corp|Corporation|Bank|Securities|Management|Holdings|Technologies|Technology|Services)\b',
            re.IGNORECASE
        )

        self.address_start_pattern = re.compile(
            r'\b(?:Building|Wing|Block|Road|Marg|Complex|East|West|Mumbai|Maharashtra|India|Pune|Bhopal|Floor|Bhavan|Nagar|Apartment|Street|Plot|Phase|Tower|Colony|Kurla|Bandra|Prabhadevi|Circle|House)\b|\b\d{6}\b|\b\d{3}\s\d{3}\b',
            re.IGNORECASE
        )

        self.address_stop_prefixes = [
            "telephone", "email", "website", "contact person", "contact name",
            "sebi", "cin:", "fax:", "tel:", "phone", "investor grievance e-mail", "investor grievance email"
        ]

        self.name_pattern = re.compile(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}\*?)\b')

        self.name_reject_keywords = {
            "limited", "ltd", "pvt", "llp", "inc", "corp", "corporation", "bank",
            "securities", "management", "road", "marg", "building", "complex",
            "wing", "block", "floor", "nagar", "apartment", "street", "plot",
            "phase", "trust", "society", "association", "university", "institute",
            "college", "school", "office", "department", "ministry", "government",
            "state", "national", "international", "global", "industries", "services",
            "solutions", "officer", "director", "secretary", "chairman", "compliance",
            "executive", "promoter", "audit", "committee", "nomination", "remuneration",
            "stakeholders", "relationship", "meeting", "board", "general", "annual",
            "extraordinary", "prospectus", "red", "herring", "issue", "offer", "bid",
            "working", "day", "days", "financial", "year", "act", "schedule", "table",
            "section", "signatory", "signatories", "january", "february", "march",
            "april", "june", "july", "august", "september", "october", "november",
            "december", "monday", "tuesday", "wednesday", "thursday", "friday",
            "saturday", "sunday"
        }

        self.fake_company_pool = [
            "Orion Capital Services Limited",
            "Maple Securities Advisory Limited",
            "Northbridge Wealth Management Limited",
            "Silverline Financial Services Limited",
            "Bluecrest Capital Markets Limited",
            "Pinnacle Securities Limited",
            "Cedarstone Investment Services Limited",
            "Summit Broking Services Limited"
        ]

        self.fake_address_pool = [
            [
                "Unit 402, Orion Business Park, Sector 62",
                "Noida 201309",
                "Uttar Pradesh, India"
            ],
            [
                "801-804, Wing B, Nexus Corporate Centre, Block C",
                "Bandra Kurla Complex, Bandra East Mumbai 400 098",
                "Maharashtra, India"
            ],
            [
                "Plot 44, Lakeview Towers, MG Road",
                "Bengaluru 560001",
                "Karnataka, India"
            ],
            [
                "H.No. 27-4, Riverfront Residency, Civil Lines",
                "Indore 452001",
                "Madhya Pradesh, India"
            ]
        ]

        self._company_map = {}
        self._address_map = {}


    def _stable_index(self, text: str, modulo: int) -> int:
        return abs(hash(text.strip().lower())) % modulo if modulo > 0 else 0

    def _get_fake_company(self, original_company_line: str) -> str:
        key = original_company_line.strip()
        if not key:
            return key
        if key in self._company_map:
            return self._company_map[key]
        idx = self._stable_index(key, len(self.fake_company_pool))
        fake_company = self.fake_company_pool[idx]
        self._company_map[key] = fake_company
        return fake_company

    def _get_fake_address_block(self, original_block_text: str, target_line_count: int):
        key = original_block_text.strip()
        if not key:
            return [""] * max(1, target_line_count)

        if key in self._address_map:
            lines = self._address_map[key]
        else:
            idx = self._stable_index(key, len(self.fake_address_pool))
            lines = list(self.fake_address_pool[idx])  # copy
            self._address_map[key] = lines

        out = list(lines)
        if len(out) < target_line_count:
            out += [""] * (target_line_count - len(out))
        elif len(out) > target_line_count:
            out = out[:target_line_count]
        return out

   
    def _redact_contact_person_line(self, text: str) -> str:
        if ":" not in text:
            return text

        prefix, names_part = text.split(":", 1)
        tokens = re.split(r'(/|,|\b(?:and|&)\b)', names_part, flags=re.IGNORECASE)

        redacted_names_part = ""
        for token in tokens:
            words = token.strip().split()
            cap_count = sum(1 for w in words if w.istitle() or w.isupper())

            if len(words) >= 2 and cap_count >= 2:
                fake_name = self.replacer.mapper.get_replacement(token.strip(), pii_type.NAME)
                left_space = token[:len(token) - len(token.lstrip())]
                right_space = token[len(token.rstrip()):]
                redacted_names_part += f"{left_space}{fake_name}{right_space}"
            else:
                redacted_names_part += token

        return f"{prefix}:{redacted_names_part}"

    def _redact_single_line(self, text: str) -> str:
        if not text.strip():
            return text

        text_lower = text.lower()

        if "contact person" in text_lower or "contact name" in text_lower:
            return self._redact_contact_person_line(text)

        is_generic_name = False
        name_match = None
        has_financial = "%" in text or "₹" in text or "rs." in text_lower or len(re.findall(r'\d', text)) > 4

        if not has_financial:
            match = self.name_pattern.search(text)
            if match:
                matched_str = match.group(1)
                words_in_match = set(matched_str.lower().split())
                contains_reject = bool(words_in_match & self.name_reject_keywords)

                if not contains_reject:
                    is_generic_name = True
                    name_match = matched_str

        allowed_types = []

        if "email" in text_lower or "e-mail" in text_lower or "@" in text_lower:
            allowed_types.append(pii_type.EMAIL)

        if "telephone" in text_lower or "phone" in text_lower or "tel:" in text_lower:
            allowed_types.append(pii_type.PHONE)


        if allowed_types:
            entries = self.detector.detect(text)
            filtered = [e for e in entries if e.entity_type in allowed_types]
            if filtered:
                text = self.replacer.redact_text(text, filtered)
            if is_generic_name and name_match:
                fake_name = self.replacer.mapper.get_replacement(name_match, pii_type.NAME)
                text = text.replace(name_match, fake_name)
            return text

        if is_generic_name and name_match:
            fake_name = self.replacer.mapper.get_replacement(name_match, pii_type.NAME)
            return text.replace(name_match, fake_name)

        return text


    def _process_paragraph_list(self, paragraphs):
        n = len(paragraphs)
        if n == 0:
            return

        redacted_indices = set()

        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue
            is_heading = (
                para.style.name.startswith("Heading")
                or (text.isupper() and len(text.split()) < 10)
            )
            if not is_heading and self.company_pattern.search(text):
                fake_company = self._get_fake_company(text)
                para.text = fake_company
                redacted_indices.add(i)

        i = 0
        while i < n:
            if i in redacted_indices:
                i += 1
                continue

            para = paragraphs[i]
            text = para.text.strip()
            word_count = len(text.split())

            if text and word_count <= 15 and self.address_start_pattern.search(text):
                block = [i]
                j = i + 1

                while j < n:
                    if j in redacted_indices:
                        break

                    nxt_text = paragraphs[j].text.strip()
                    nxt_lower = nxt_text.lower()

                    if not nxt_text:
                        break
                    if len(nxt_text.split()) > 15:
                        break
                    if any(nxt_lower.startswith(p) for p in self.address_stop_prefixes):
                        break

                    block.append(j)
                    j += 1

                orig_block_text = "\n".join(paragraphs[idx].text for idx in block)
                fake_lines = self._get_fake_address_block(orig_block_text, len(block))

                for line_idx, para_idx in enumerate(block):
                    paragraphs[para_idx].text = fake_lines[line_idx]
                    redacted_indices.add(para_idx)

                i = j
            else:
                i += 1

        for i, para in enumerate(paragraphs):
            if i in redacted_indices:
                continue
            para.text = self._redact_single_line(para.text)

    def process_and_save(self):
        doc = Document(self.input_path)

        self._process_paragraph_list(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_paragraph_list(cell.paragraphs)

        for section in doc.sections:
            self._process_paragraph_list(section.header.paragraphs)
            self._process_paragraph_list(section.footer.paragraphs)

        doc.save(self.output_path)
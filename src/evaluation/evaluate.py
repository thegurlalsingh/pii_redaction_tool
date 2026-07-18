import re
from collections import defaultdict
from docx import Document

from detectors.hybrid_detector import hybrid_detector
from data_structure import pii_type
import os


def read_docx_text(path: str) -> str:
    doc = Document(path)
    chunks = []

    for p in doc.paragraphs:
        if p.text:
            chunks.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        chunks.append(p.text)

    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text:
                chunks.append(p.text)
        for p in section.footer.paragraphs:
            if p.text:
                chunks.append(p.text)

    return "\n".join(chunks)


def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip().lower()


def metrics(tp: int, fp: int, fn: int):
    precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
    recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
    accuracy = tp / (tp + fp + fn) if (tp + fp + fn) > 0 else 0.0
    return precision, recall, accuracy


def main():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    original_docx = os.path.join(base_dir, "data", "input", "Red Herring Prospectus.docx")
    redacted_docx = os.path.join(base_dir, "data", "output", "redacted.docx")
    report_out = os.path.join(base_dir, "data", "reports", "redacted.docx")

    original_text = read_docx_text(original_docx)
    redacted_text = read_docx_text(redacted_docx)

    detector = hybrid_detector()

    orig_entities = detector.detect(original_text)
    red_entities = detector.detect(redacted_text)

    orig_by_type = defaultdict(set)
    red_by_type = defaultdict(set)

    for e in orig_entities:
        t = normalize(e.text)
        if t:
            orig_by_type[e.entity_type].add(t)

    for e in red_entities:
        t = normalize(e.text)
        if t:
            red_by_type[e.entity_type].add(t)

    rows = []
    total_tp = total_fp = total_fn = 0

    for ptype in pii_type:
        orig_set = orig_by_type.get(ptype, set())
        red_set = red_by_type.get(ptype, set())

        still_present = orig_set.intersection(red_set)
        removed = orig_set - red_set

        tp = len(removed)
        fn = len(still_present)

        fp = len(red_set - orig_set)

        if tp + fp + fn == 0:
            continue

        total_tp += tp
        total_fp += fp
        total_fn += fn

        p, r, a = metrics(tp, fp, fn)
        rows.append((ptype.value, tp, fp, fn, p, r, a))

    o_p, o_r, o_a = metrics(total_tp, total_fp, total_fn)

    lines = []
    
    lines.append("")
    lines.append(f"{'PII Type':<18} | {'TP':>4} | {'FP':>4} | {'FN':>4} | {'Precision':>9} | {'Recall':>9} | {'Accuracy':>9}")
    lines.append("-" * 85)

    for name, tp, fp, fn, p, r, a in rows:
        lines.append(f"{name:<18} | {tp:>4} | {fp:>4} | {fn:>4} | {p:>9.2f} | {r:>9.2f} | {a:>9.2f}")

    lines.append("-" * 85)
    lines.append(f"{'OVERALL':<18} | {total_tp:>4} | {total_fp:>4} | {total_fn:>4} | {o_p:>9.2f} | {o_r:>9.2f} | {o_a:>9.2f}")

    report = "\n".join(lines)
    print(report)

    with open(report_out, "w", encoding="utf-8") as f:
        f.write(report)

    print(f"\nSaved report to: {report_out}")


if __name__ == "__main__":
    main()